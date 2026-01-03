import sqlite3
import os
import docx
import textract
from bs4 import BeautifulSoup
from datetime import datetime

# Database Configuration
DB_NAME = "database.db"

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS projects (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    name TEXT UNIQUE NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )''')
    c.execute('''CREATE TABLE IF NOT EXISTS modules (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    project_id INTEGER,
                    name TEXT,
                    FOREIGN KEY(project_id) REFERENCES projects(id),
                    UNIQUE(project_id, name)
                )''')
    c.execute('''CREATE TABLE IF NOT EXISTS test_cases (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    module_id INTEGER,
                    content TEXT,
                    status TEXT DEFAULT 'PENDING',
                    bug_report TEXT,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY(module_id) REFERENCES modules(id)
                )''')
    conn.commit()
    conn.close()

def get_db_connection():
    conn = sqlite3.connect(DB_NAME)
    conn.row_factory = sqlite3.Row
    return conn

# --- Project Management ---
def get_all_projects():
    conn = get_db_connection()
    try:
        # Try with created_at first
        rows = conn.execute("SELECT name, created_at FROM projects ORDER BY id DESC").fetchall()
        result = [{"name": row["name"], "created_at": row["created_at"]} for row in rows]
    except:
        # Fallback for old databases without created_at
        rows = conn.execute("SELECT name FROM projects ORDER BY id DESC").fetchall()
        result = [{"name": row["name"], "created_at": None} for row in rows]
    conn.close()
    return result

def create_project(name):
    conn = get_db_connection()
    try:
        conn.execute("INSERT INTO projects (name) VALUES (?)", (name,))
        conn.commit()
        return True
    except sqlite3.IntegrityError:
        return False  # Project already exists
    finally:
        conn.close()

def delete_project(name):
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT id FROM projects WHERE name = ?", (name,))
    proj = c.fetchone()
    if not proj:
        conn.close()
        return False
    
    proj_id = proj['id']
    
    # Cascade delete: test_cases -> modules -> project
    c.execute("DELETE FROM test_cases WHERE module_id IN (SELECT id FROM modules WHERE project_id = ?)", (proj_id,))
    c.execute("DELETE FROM modules WHERE project_id = ?", (proj_id,))
    c.execute("DELETE FROM projects WHERE id = ?", (proj_id,))
    
    # Reset auto-increment if no cases left
    c.execute("SELECT COUNT(*) FROM test_cases")
    if c.fetchone()[0] == 0:
        c.execute("DELETE FROM sqlite_sequence WHERE name='test_cases'")
    
    conn.commit()
    conn.close()
    return True

def get_project_stats(project_name):
    """Get statistics for a project"""
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT id FROM projects WHERE name = ?", (project_name,))
    proj = c.fetchone()
    if not proj:
        conn.close()
        return None
    
    proj_id = proj['id']
    
    stats = conn.execute("""
        SELECT 
            COUNT(t.id) as total,
            SUM(CASE WHEN t.status = 'Pass' THEN 1 ELSE 0 END) as passed,
            SUM(CASE WHEN t.status = 'FAILED' THEN 1 ELSE 0 END) as failed,
            SUM(CASE WHEN t.status = 'PENDING' THEN 1 ELSE 0 END) as pending
        FROM test_cases t
        JOIN modules m ON t.module_id = m.id
        WHERE m.project_id = ?
    """, (proj_id,)).fetchone()
    
    module_count = conn.execute("""
        SELECT COUNT(*) FROM modules WHERE project_id = ?
    """, (proj_id,)).fetchone()[0]
    
    conn.close()
    
    return {
        "total_cases": stats["total"] or 0,
        "passed": stats["passed"] or 0,
        "failed": stats["failed"] or 0,
        "pending": stats["pending"] or 0,
        "modules": module_count
    }

# --- Document Reading ---
def read_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    # Read Paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            full_text.append(paragraph.text)
    
    # Read Tables (crucial for requirements in tables)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                full_text.append(" | ".join(row_text))
                
    return "\n".join(full_text)

def read_doc(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        if "<html" in content or "MIME-Version" in content:
            soup = BeautifulSoup(content, 'html.parser')
            for script in soup(["script", "style", "meta", "link", "xml"]):
                script.decompose()
            text = soup.get_text(separator='\n')
            return "\n".join([line.strip() for line in text.splitlines() if line.strip()])
        text = textract.process(file_path).decode('utf-8')
        return text
    except Exception as e:
        if "antiword" in str(e):
            return "Format Error: Old .doc file. Please save as .docx."
        raise e

def read_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except:
        with open(file_path, 'r', encoding='cp1251') as f:
            return f.read()

# --- Database Operations ---
def add_cases(cases_list, module_name, project_name="togetherfun"):
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT id FROM projects WHERE name = ?", (project_name,))
    project = c.fetchone()
    if not project:
        c.execute("INSERT INTO projects (name) VALUES (?)", (project_name,))
        project_id = c.lastrowid
    else:
        project_id = project['id']
    
    c.execute("SELECT id FROM modules WHERE project_id = ? AND name = ?", (project_id, module_name))
    module = c.fetchone()
    if not module:
        c.execute("INSERT INTO modules (project_id, name) VALUES (?, ?)", (project_id, module_name))
        module_id = c.lastrowid
    else:
        module_id = module['id']
    
    for case in cases_list:
        content = case
        
        # Check if case is a string that looks like a dict/JSON
        if isinstance(case, str) and case.strip().startswith('{') and case.strip().endswith('}'):
            try:
                import ast
                # ast.literal_eval is safer than eval for stringified dicts
                parsed = ast.literal_eval(case)
                if isinstance(parsed, dict):
                    case = parsed
            except:
                pass

        if isinstance(case, dict):
            # Try various keys the AI might use
            steps = case.get("steps") or case.get("description") or case.get("content") or case.get("text")
            result = case.get("result") or case.get("expected_result") or case.get("expected")
            
            if steps and result:
                content = f"Кроки: {steps} <br> Очікуваний результат: {result}"
            elif steps:
                content = steps
            else:
                content = str(case)
        
        c.execute("INSERT INTO test_cases (module_id, content, status) VALUES (?, ?, 'PENDING')",
                  (module_id, content))
    conn.commit()
    conn.close()

def get_unique_pending_modules(project_name="togetherfun"):
    conn = get_db_connection()
    query = """
        SELECT m.name, MIN(t.id) as first_case_id
        FROM modules m
        JOIN test_cases t ON m.id = t.module_id
        JOIN projects p ON m.project_id = p.id
        WHERE p.name = ? AND t.status = 'PENDING'
        GROUP BY m.name
        ORDER BY m.id DESC
    """
    rows = conn.execute(query, (project_name,)).fetchall()
    conn.close()
    return {row['name']: row['first_case_id'] for row in rows}

def get_module_stats(project_name="togetherfun"):
    """Get statistics for each module: total cases, passed, failed, pending"""
    conn = get_db_connection()
    query = """
        SELECT 
            m.name,
            COUNT(t.id) as total,
            SUM(CASE WHEN t.status = 'Pass' THEN 1 ELSE 0 END) as passed,
            SUM(CASE WHEN t.status = 'FAILED' THEN 1 ELSE 0 END) as failed,
            SUM(CASE WHEN t.status = 'PENDING' THEN 1 ELSE 0 END) as pending
        FROM modules m
        JOIN test_cases t ON m.id = t.module_id
        JOIN projects p ON m.project_id = p.id
        WHERE p.name = ?
        GROUP BY m.name
        ORDER BY m.id DESC
    """
    rows = conn.execute(query, (project_name,)).fetchall()
    conn.close()
    return [{
        'name': row['name'],
        'total': row['total'],
        'passed': row['passed'],
        'failed': row['failed'],
        'pending': row['pending'],
        'progress': round((row['passed'] / row['total']) * 100) if row['total'] > 0 else 0
    } for row in rows]


def get_next_pending_case_by_module(module_name, project_name="togetherfun"):
    conn = get_db_connection()
    
    # First try to get PENDING cases
    query_pending = """
        SELECT t.id, t.content, t.status
        FROM test_cases t
        JOIN modules m ON t.module_id = m.id
        JOIN projects p ON m.project_id = p.id
        WHERE m.name = ? AND p.name = ? AND t.status = 'PENDING'
        ORDER BY t.id ASC
        LIMIT 1
    """
    row = conn.execute(query_pending, (module_name, project_name)).fetchone()
    
    # If no pending cases, try to get FAILED cases for retesting
    if not row:
        query_failed = """
            SELECT t.id, t.content, t.status
            FROM test_cases t
            JOIN modules m ON t.module_id = m.id
            JOIN projects p ON m.project_id = p.id
            WHERE m.name = ? AND p.name = ? AND t.status = 'FAILED'
            ORDER BY t.id ASC
            LIMIT 1
        """
        row = conn.execute(query_failed, (module_name, project_name)).fetchone()
    
    conn.close()
    if row: 
        return {
            "id": row['id'], 
            "text": row['content'],
            "status": row['status'],
            "is_retest": row['status'] == 'FAILED'
        }
    return None

def update_case_status(case_id, status, bug_report=None):
    conn = get_db_connection()
    if bug_report:
        conn.execute("UPDATE test_cases SET status = ?, bug_report = ? WHERE id = ?", (status, bug_report, case_id))
    else:
        conn.execute("UPDATE test_cases SET status = ? WHERE id = ?", (status, case_id))
    conn.commit()
    conn.close()

def get_failed_cases_with_bugs(project_name="togetherfun"):
    conn = get_db_connection()
    query = """
        SELECT t.id, m.name as module_name, t.content, t.bug_report
        FROM test_cases t
        JOIN modules m ON t.module_id = m.id
        JOIN projects p ON m.project_id = p.id
        WHERE p.name = ? AND t.status = 'FAILED'
        ORDER BY t.id DESC
    """
    rows = conn.execute(query, (project_name,)).fetchall()
    conn.close()
    return [{
        "id": row['id'],
        "module": row['module_name'],
        "case_text": row['content'],
        "bug_report": row['bug_report']
    } for row in rows]

def update_bug_report_text(case_id, new_text):
    conn = get_db_connection()
    conn.execute("UPDATE test_cases SET bug_report = ? WHERE id = ?", (new_text, case_id))
    conn.commit()
    conn.close()

def delete_bug_report(case_id):
    conn = get_db_connection()
    conn.execute("DELETE FROM test_cases WHERE id = ?", (case_id,))
    conn.commit()
    conn.close()

# --- Bulk & Pagination Helper ---
def get_all_cases_paginated(project_name="togetherfun", page=1, limit=20, status=None):
    offset = (page - 1) * limit
    conn = get_db_connection()
    
    # Check if project exists
    c = conn.cursor()
    c.execute("SELECT id FROM projects WHERE name = ?", (project_name,))
    project = c.fetchone()
    if not project:
        conn.close()
        return [], 0
    proj_id = project['id']
    
    where_clause = "WHERE m.project_id = ?"
    params = [proj_id]
    
    if status and status != 'all':
        where_clause += " AND t.status = ?"
        params.append(status)

    # Total Count
    count_query = f"""
        SELECT COUNT(*) as total
        FROM test_cases t
        JOIN modules m ON t.module_id = m.id
        {where_clause}
    """
    total = conn.execute(count_query, params).fetchone()['total']

    # Items
    query = f"""
        SELECT t.id, m.name as module, t.content, t.status, t.bug_report
        FROM test_cases t
        JOIN modules m ON t.module_id = m.id
        {where_clause}
        ORDER BY t.id DESC
        LIMIT ? OFFSET ?
    """
    params.extend([limit, offset])
    rows = conn.execute(query, params).fetchall()
    
    # Get all unique modules for this project for the filter
    modules_query = "SELECT DISTINCT name FROM modules WHERE project_id = ?"
    module_rows = conn.execute(modules_query, (proj_id,)).fetchall()
    all_modules = [r['name'] for r in module_rows]
    
    print(f"DEBUG: Project {project_name} (ID {proj_id}) - Fetching cases offset {offset} limit {limit}. Found: {len(rows)}")
    conn.close()
    
    return [dict(row) for row in rows], total, all_modules

def delete_cases_bulk(case_ids):
    conn = get_db_connection()
    placeholders = ','.join(['?'] * len(case_ids))
    sql = f"DELETE FROM test_cases WHERE id IN ({placeholders})"
    conn.execute(sql, tuple(case_ids))
    conn.commit()
    conn.close()

def update_cases_status_bulk(case_ids, status):
    conn = get_db_connection()
    placeholders = ','.join(['?'] * len(case_ids))
    sql = f"UPDATE test_cases SET status = ? WHERE id IN ({placeholders})"
    args = [status] + case_ids
    conn.execute(sql, tuple(args))
    conn.commit()
    conn.close()

def delete_all_cases_for_project(project_name):
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT id FROM projects WHERE name = ?", (project_name,))
    proj = c.fetchone()
    if not proj:
        conn.close()
        return
    proj_id = proj['id']
    query = "DELETE FROM test_cases WHERE module_id IN (SELECT id FROM modules WHERE project_id = ?)"
    conn.execute(query, (proj_id,))
    conn.execute("DELETE FROM modules WHERE project_id = ?", (proj_id,))
    conn.execute("DELETE FROM sqlite_sequence WHERE name='test_cases'")
    conn.commit()
    conn.close()

def reset_module_cases(project_name, module_name):
    """Resets all cases in a module to PENDING and clears bug reports"""
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT id FROM projects WHERE name = ?", (project_name,))
    project = c.fetchone()
    if not project:
        conn.close()
        return False
    c.execute("SELECT id FROM modules WHERE project_id = ? AND name = ?", (project['id'], module_name))
    module = c.fetchone()
    if not module:
        conn.close()
        return False
    c.execute("UPDATE test_cases SET status = 'PENDING', bug_report = NULL WHERE module_id = ?", (module['id'],))
    conn.commit()
    conn.close()
    return True